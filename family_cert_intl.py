import os, re, zipfile, unicodedata
from io import BytesIO
from datetime import datetime
from dotenv import load_dotenv


import boto3
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE



def set_cell_vertical(cell, dir="btLr"):
    """
    Rotate text in this cell bottom-to-top, left-to-right.
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    td   = OxmlElement('w:textDirection')
    td.set(qn('w:val'), dir)
    tcPr.append(td)

# ── AWS / ENV ───────────────────────────────────────────────────────────────
load_dotenv()
textract = boto3.client(
    "textract",
    aws_access_key_id     = os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key = os.getenv("AWS_SECRET_ACCESS_KEY"),
    region_name           = os.getenv("AWS_REGION") or "us-east-2"
)

# ── STREAMLIT UI ────────────────────────────────────────────────────────────
st.set_page_config(page_title="AI Translator - Certifikata Familjare", layout="centered")
st.title("Certifikata Familjare: Shqip  -  Italisht")
st.markdown("Upload one or more family certificates (PDF or image), then download the Italian DOCX.")

uploaded_files = st.file_uploader(
    "Upload certificate(s)",
    type=["pdf", "jpg", "jpeg", "png"],
    accept_multiple_files=True
)

# --- Simple password gate (one shared password) ---
password = st.text_input("Password", type="password")
if "APP_PASSWORD" not in st.secrets:
    st.stop()  # safety if not configured
if password != st.secrets["APP_PASSWORD"]:
    st.warning("Enter the password to continue")
    st.stop()
# ---------------------------------------------------

# ── HELPER: TEXTRACT OCR WRAPPER (with page tagging) ─────────────────────────
from PIL import Image

def correct_orientation(img: Image.Image) -> Image.Image:
    """
    If the image is taller than it is wide, rotate it
    so that it’s landscape (width>=height).
    """
    if img.height > img.width:
        return img.rotate(-90, expand=True)
    return img

def get_textract_blocks(uploaded_file):
    from botocore.exceptions import ClientError
    from pdf2image import convert_from_bytes
    from io import BytesIO

    data = uploaded_file.read()
    name = uploaded_file.name.lower()

    def analyze_bytes(bts, page_no=None):
        blocks = textract.analyze_document(
            Document={'Bytes': bts},
            FeatureTypes=["TABLES", "FORMS"]
        )["Blocks"]
        if page_no is not None:
            for b in blocks:
                b["Page"] = page_no
        return blocks

    # PDF path
    if name.endswith(".pdf"):
        try:
            # Try native PDF
            return analyze_bytes(data)
        except ClientError as e:
            if "UnsupportedDocumentException" in str(e):
                # Fallback → image conversion + orientation fix
                all_blocks = []
                pages = convert_from_bytes(data, dpi=300)
                for idx, pil_img in enumerate(pages, start=1):
                    fixed = correct_orientation(pil_img)
                    buf = BytesIO()
                    fixed.save(buf, format="PNG")
                    buf.seek(0)
                    all_blocks.extend(analyze_bytes(buf.getvalue(), page_no=idx))
                return all_blocks
            else:
                raise

    # Image path (JPG/PNG etc)
    img = Image.open(BytesIO(data))
    fixed = correct_orientation(img)
    buf = BytesIO()
    fixed.save(buf, format="PNG")
    buf.seek(0)
    return analyze_bytes(buf.getvalue(), page_no=1)

# -- HELPER: translator

def _norm(s: str) -> str:
    if not s: return ""
    # lower, strip, remove accents, collapse spaces and punctuation/slashes
    s = s.strip().lower()
    s = ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')
    s = re.sub(r'[\s\.\-_/]+', ' ', s).strip()
    return s

RELATION_MAP = {
    "kryefamiljar": "Capofamiglia",
    "i biri": "Figlio",
    "biri": "Figlio",
    "e bija": "Figlia",
    "bija": "Figlia",
    "bashkeshortja": "Moglie",
    "bashkeshorti": "Marito",
    "gruaja": "Moglie",
    "burri": "Marito",
    "nipi": "Nipote (maschio)",
    "mbesa": "Nipote (femmina)",
    "babai": "Padre",
    "nena": "Madre",
    "gjyshi": "Nonno",
    "gjyshja": "Nonna",
    "vellai": "Fratello",
    "motra": "Sorella"
}

MARITAL_MAP = {
    "i martuar":  ("Coniugato", "Coniugata"),
    "e martuar":  ("Coniugato", "Coniugata"),
    "i/e martuar":("Coniugato", "Coniugata"),
    "beqar":      ("Celibe", "Nubile"),
    "beqare":     ("Celibe", "Nubile"),
    "beqar/e":    ("Celibe", "Nubile"),
    "i/e ve":     ("Vedovo", "Vedova"),
    "i ve":       ("Vedovo", "Vedova"),
    "e ve":       ("Vedovo", "Vedova"),
    "i/e divorcuar": ("Divorziato", "Divorziata"),
    "i/e ndare":     ("Separato", "Separata")
}

CITIZENSHIP_MAP = {
    # what you asked for:
    "shqiptare": "Albanese",
    "shqiptar":  "Albanese",
    "shqiptar/e":"Albanese",

    # a few examples to extend (optional):
    "italian":      ("Italiano", "Italiana"),
    "italiane":     ("Italiano", "Italiana"),
    "grek":         ("Greco", "Greca"),
    "greke":        ("Greco", "Greca"),
    "francez":      "Francese",
    "franceze":     "Francese",
}

def translate_citizenship(alb_cit: str, sex: str) -> str:
    n = _norm(alb_cit)
    if n in CITIZENSHIP_MAP:
        it = CITIZENSHIP_MAP[n]
        if isinstance(it, tuple):
            return it[1] if (sex or "").strip().upper() == "F" else it[0]
        return it
    return alb_cit

def translate_relation(alb_relation: str, sex: str) -> str:
    n = _norm(alb_relation)
    if n in RELATION_MAP:
        return RELATION_MAP[n]
    return alb_relation

def translate_marital_status(alb_status: str, sex: str) -> str:
    n = _norm(alb_status)
    s = (sex or "").strip().upper()
    if n in MARITAL_MAP:
        male, female = MARITAL_MAP[n]
        return female if s == "F" else male
    return alb_status



# ── HELPER: Extract Issue Date (dd.MM.yyyy) ─────────────────────────────────
def extract_issue_date(blocks):
    import re
    for b in blocks:
        if b["BlockType"] == "LINE":
            txt = b["Text"].strip()
            if re.fullmatch(r"\d{1,2}\.\d{2}\.\d{4}", txt):
                return txt
    return ""

# ── HELPER: SEAL FOOTER ──────────────────────────────────────────────────────
def extract_seal_footer(blocks):
    import re

    # 1) collect every LINE
    lines = [b["Text"].strip() for b in blocks if b["BlockType"]=="LINE"]

    # 2) find the last “vulosur elektronikisht…”
    idxs = [i for i, txt in enumerate(lines)
            if "vulosur elektronikisht" in txt.lower()]
    if not idxs:
        return ""
    start = idxs[-1]

    # 3) stitch together the next 6 lines into one snippet
    snippet = "\n".join(lines[start : start + 6])

    # 4) pull the first date-like thing (yyyy/mm/dd …)
    date_line = ""
    m = re.search(r"\b\d{4}/\d{2}/\d{2}.*", snippet)
    if m:
        cleaned = re.sub(r"^(Date|Datë|Daté)\s*:?\s*", "",
                         m.group(0), flags=re.I).strip()
        date_line = f"In data: {cleaned}"

    # 5) harvest _all_ 30–40 char runs (allowing OCR’d “O” too)
    raw = re.findall(r"\b[0-9A-Fa-fO]{30,40}\b", snippet)
    # normalize any leading O → 0, then keep only true hex
    cands = [
        re.sub(r'^[Oo]', '0', h) 
        for h in raw
    ]
    cands = [h for h in cands if re.fullmatch(r"[0-9A-Fa-f]{30,40}", h)]
    hash_line = max(cands, key=len) if cands else ""

    # 6) if we got nothing, bail
    if not (date_line or hash_line):
        return ""

    # 7) build your final 4-line seal text
    return "\n".join([
        "Timbrato elettronicamente dalla Direzione",
        "Generale dello Stato Civile",
        date_line,
        hash_line
    ])




# ── TABLE-FIELD EXTRACTION ON PAGE 2 ────────────────────────────────────────
def extract_family_table_v2(blocks, bmap):
    import re

    # 1) find the TABLE on page 2
    tbl = next((b for b in blocks 
                if b["BlockType"]=="TABLE" and b.get("Page")==2), None)
    if not tbl:
        return {"header": [], "rows": [], "seal_footer": ""}

    # 2) build a map: rows_map[row_index][col_index] = cell_text
    rows_map = {}
    for rel in tbl.get("Relationships", []):
        if rel["Type"]!="CHILD": 
            continue
        for cid in rel["Ids"]:
            cell = bmap[cid]
            if cell["BlockType"]!="CELL": 
                continue
            r, c = cell["RowIndex"], cell["ColumnIndex"]
            txt = " ".join(
                bmap[wid]["Text"]
                for ch in cell.get("Relationships", [])
                if ch["Type"]=="CHILD"
                for wid in ch["Ids"]
                if bmap[wid]["BlockType"]=="WORD"
            ).strip()
            rows_map.setdefault(r, {})[c] = txt

    # 3) extract header row (row 1)
    max_col = max((c for row in rows_map.values() for c in row.keys()), default=0)
    header = [ rows_map.get(1, {}).get(c, "") for c in range(1, max_col+1) ]

    # 4) extract data rows 3–12
    data_rows = []
for idx, r in enumerate(range(3, 13), start=1):
    row = rows_map.get(r, {})
    sex = (row.get(5, "") or "").strip().upper()  # "M" / "F"
    dob = "/".join(filter(None, [row.get(7,""), row.get(8,""), row.get(9,"")]))

    # Apply glossary/translation
    rel_it  = translate_relation(row.get(6, ""), sex)
    stat_it = translate_marital_status(row.get(10, ""), sex)
    cit_it   = translate_citizenship(row.get(12, ""),  sex)

    data_rows.append({
        "N.":                            str(idx),
        "1. Nome e Cognome":             row.get(2, ""),
        "2. Nome del Padre":             row.get(3, ""),
        "3. Nome della Madre":           row.get(4, ""),
        "4. Sesso":                      sex,       # ← use normalized sex
        "5. Legame con il capofamiglia": rel_it,    # ← translated
        "6. Data di nascita":            dob,
        "7. Stato Civile":               stat_it,   # ← translated
        "8. Luogo di Nascita":           row.get(11, ""),
        "9. Cittadinanza":               cit_it,    # translated
        "10. Numero Personale":          row.get(13, ""),
    })

    # 5) extract the seal footer from *all* lines
    seal_footer = extract_seal_footer(blocks)

return {
    "header":      header,
    "rows":        data_rows,
    "seal_footer": seal_footer
}


# ── HEADER (Comune / Sezione) ───────────────────────────────────────────────
def extract_comune_sezione(blocks):
    import re
    # don’t filter by Page—Textract images have no Page attribute
    lines = [b["Text"].strip() for b in blocks
             if b["BlockType"] == "LINE"]

    comune = sezione = ""
    for i, ln in enumerate(lines):
        if "Bashkia" in ln:
            m = re.search(r"Bashkia\s+([A-ZÇËA-Za-zë\-]+)", ln)
            if m:
                comune = m.group(1).title()
        if "Njësia Administrative" in ln or "Njesia Administrative" in ln:
            parts = ln.split("Administrative", 1)
            if len(parts) > 1:
                s = parts[1].strip()
                # if the “nr.” is alone, pull the next line too
                if s.lower() in ("nr.", "nr"):
                    s += " " + (lines[i + 1] if i + 1 < len(lines) else "")
                sezione = s.title()
    return comune, sezione


# ── DOCX TEMPLATE ───────────────────────────────────────────────────────────
def make_docx(people, comune, sezione, seal_text):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Mm
    from docx.oxml.ns import qn
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from datetime import datetime
    import os
    from io import BytesIO

    doc = Document()
    today = datetime.today().strftime("%d.%m.%Y")

    # A4 landscape
    sec = doc.sections[0]
    sec.orientation   = WD_ORIENT.LANDSCAPE
    sec.page_width    = Mm(297)
    sec.page_height   = Mm(210)
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(1)
    sec.left_margin   = Cm(2)
    sec.right_margin  = Cm(2)

    # base style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0,0,0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    def addp(txt, size=11, align="left", bold=False, italic=False, underline=False, indent=0):
        p = doc.add_paragraph()
        r = p.add_run(txt)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)
        r.bold, r.italic, r.underline = bold, italic, underline
        r.font.color.rgb = RGBColor(0,0,0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = 1
        if indent: p.paragraph_format.left_indent = Cm(indent)
        p.alignment = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }[align]
        return p

    # header (flag + Comune/Sezione)
    hdr = doc.add_table(1,2); hdr.style="Table Grid"; hdr.autofit=True
    c1 = hdr.cell(0,0).paragraphs[0]
    img = os.path.join(os.getcwd(),"al_flag.png")
    if os.path.exists(img):
        run = c1.add_run(); run.add_picture(img, width=Cm(0.9))
    c1.add_run("\n\nREPUBBLICA D'ALBANIA\n\n").bold = True
    c1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    c2 = hdr.cell(0,1).paragraphs[0]
    parts = []
    if comune:   parts.append(f"\n\n\nUfficio di Stato Civile Comune di {comune}")
    if sezione: parts.append(f"Sezione Amministrativa {sezione}")
    c2.add_run("\n".join(parts)).bold = True
    c2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # ── Issue Date under header ─────────────────────────────────────────────
    # this paragraph will sit just below the header table

    p_date = doc.add_paragraph(issue_date)
    p_date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_date.paragraph_format.first_line_indent = Cm(23)
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after  = Pt(6)   # small gap before next content

    # ── Then your title, data table, seal, footer etc. continue here …
    # title
    p0 = doc.add_paragraph()
    run = p0.add_run("\nCERTIFICATO DI STATO DI FAMIGLIA\n")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p0.paragraph_format.line_spacing = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
   


    p0 = doc.add_paragraph()
    run = p0.add_run("\n\nIn base al Registro Nazionale dello Stato Civile dell’anno 2010, si certificano i seguenti dati:\n")
    run.bold = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p0.paragraph_format.line_spacing = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    

    addp("\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\nTraduzione eseguita da:\nVjollca META", size=11, align="center", indent=18)

    doc.add_page_break()

    # data table with 11 columns
    cols = [
        "N.",
        "1. Nome e Cognome",
        "2. Nome del Padre",
        "3. Nome della Madre",
        "4. Sesso",
        "5. Legame con il capofamiglia",
        "6. Data di nascita",
        "7. Stato Civile",
        "8. Luogo di Nascita",
        "9. Cittadinanza",
        "10. Numero Personale"
    ]
    dt = doc.add_table(1, len(cols))
    dt.style = "Table Grid"


    dt.autofit = False

    # 2) force Word into fixed-layout mode
    tbl = dt._tbl  # low-level CT_Tbl element
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    tbl_layout = OxmlElement('w:tblLayout')
    tbl_layout.set(qn('w:type'), 'fixed')
    tbl_pr.append(tbl_layout)

    # 3) now set your exact column widths
    widths_cm = [0.9, 4, 2.4, 2.4, 0.7, 3, 2.4, 2.5, 4, 1.9, 2.7]
    for col, w in zip(dt.columns, widths_cm):
        col.width = Cm(w)
        for cell in col.cells:
            cell.width = Cm(w)

    for i,h in enumerate(cols):
        cell = dt.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for cell in dt.rows[0].cells:
        txt = cell.text.strip()
        if txt in ("4. Sesso", "9. Cittadinanza"):
            set_cell_vertical(cell, dir="btLr")

    hdr_row = dt.rows[0]
    hdr_row.height = Cm(3)                            # make header row 2 cm tall
    hdr_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY # enforce exact height

    # data rows
    for person in people:
        cells = dt.add_row().cells
        for i,key in enumerate(cols):
            val = person.get(key,"")
            cells[i].text = val
            cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for row in dt.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
  
    
    addp("Questo certificato viene rilasciato per uso all’estero", size=11, align="left")

    addp("\nTimbrato elettronicamente dalla Direzione Generale dello Stato Civile", size=11, align="right")


    if seal_text:
# split & drop empties
        raw = [l.strip() for l in seal_text.splitlines() if l.strip()]

        # grab the date line
        date_src = next((l for l in raw if re.search(r'\d{4}/\d{2}/\d{2}', l)), "")
        cleaned_date = re.sub(r'^(Date|Datë|Daté)\s*:?', '', date_src, flags=re.I).strip()
        date_line = f"In data: {cleaned_date}"

        # search for a hex string of length 30–40 anywhere in any line
        hash_match = None
        for txt in raw:
            # remove everything except hex digits
            candidate = re.sub(r'[^0-9A-Fa-f]', '', txt)
            # if it’s 30–40 chars long, that’s our hash
            if 30 <= len(candidate) <= 40:
                hash_match = candidate
                break

        # build exactly the four lines you want
        seal_lines = [
            "Timbrato elettronicamente dalla Direzione",
            "Generale dello Stato Civile",
            date_line
        ]
        if hash_match:
            seal_lines.append(hash_match)

        # 4) render them as 10 pt, zero spacing
    for line in seal_lines:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.italic = True
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # translator footnote
    addp(
        "\nNota: Questo documento è stato generato e timbrato\nda una procedura automatica da un sistema elettronico\n"
        "(Direzione Generale di Stato Civile)\n",
        size=10, italic=True
    )
    cert = (
        "Io, Vjollca META, traduttrice ufficiale della lingua italiana certificata dal Ministero "
        "della Giustizia con il numero di certificato 412 datato 31.07.2024, dichiaro di aver tradotto "
        "il testo presentatomi dalla lingua albanese all'italiano con precisione e responsabilità legale.\n"
        f"In data {today}."
    )
    tbl = doc.add_table(1,1); tbl.style="Table Grid"; tbl.autofit=False
    tbl.columns[0].width = Cm(11); tbl.rows[0].cells[0].width = Cm(13.5)
    p = tbl.cell(0,0).paragraphs[0]
    p.add_run(cert).font.size = Pt(10)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    addp("\nTraduzione eseguita da:\nVjollca META", size=11, align="center", indent=18)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf



# ── MAIN FLOW ───────────────────────────────────────────────────────────────
if uploaded_files and st.button("Translate"):
    single = len(uploaded_files)==1
    if single:
        files = [uploaded_files[0]]
    else:
        files = uploaded_files

    zip_buf = BytesIO() if not single else None
    if not single:
        zipf = zipfile.ZipFile(zip_buf, "w")

    for f in files:
        with st.spinner(f"Processing {f.name}..."):
            blocks = get_textract_blocks(f)
            bmap   = {b["Id"]: b for b in blocks}
            table_data = extract_family_table_v2(blocks, bmap)
            seal_text = extract_seal_footer(blocks)
            headers    = table_data["header"]   # if you need to render them dynamically
            people     = table_data["rows"]
            comune,sez = extract_comune_sezione(blocks)
            seal   = extract_seal_footer(blocks)
            issue_date = extract_issue_date(blocks)  
            docx_b = make_docx(people, comune, sez, seal)

        if single:
            name = f"Certificato_di_Famiglia_{datetime.today():%d-%m-%Y}.docx"
            st.download_button("📥 Download DOCX", docx_b.getvalue(), file_name=name,
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            out_name = f"{os.path.splitext(f.name)[0]}_{datetime.today():%d-%m-%Y}.docx"
            zipf.writestr(out_name, docx_b.getvalue())

    if not single:
        zipf.close()
        zip_buf.seek(0)
        st.download_button("📥 Download All Translations (ZIP)", zip_buf,
                           file_name=f"certificati_tradotti_{datetime.today():%Y-%m-%d}.zip",
                           mime="application/zip")
